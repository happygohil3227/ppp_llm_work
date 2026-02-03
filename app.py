import json
import os
import re
import time
from pathlib import Path
from typing import Dict, List

import boto3
import streamlit as st
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_openai import ChatOpenAI

from utils.vecDb import LegalVectorDB


BASE_DIR = Path(__file__).resolve().parent
TEXT_DIR = BASE_DIR / "extracted_text"
VECTOR_DB_DIR = BASE_DIR / "vectorDB"
OUTPUT_DIR = BASE_DIR / "outputs"

TEXT_DIR.mkdir(exist_ok=True)
VECTOR_DB_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)


SECTIONS: Dict[int, Dict[str, str]] = {
    1: {
        "name": "Context & Objective",
        "query": (
            "Extract the background, policy rationale, and purpose of this PPP "
            "concession agreement, including the problem it seeks to address, "
            "the public interest objective, the scope of the concession, and "
            "the intended outcomes for infrastructure service delivery and "
            "market efficiency"
        ),
    },
    2: {
        "name": "Scope of Concession & Rights Granted",
        "query": (
            "Extract clauses defining the scope of the concession, activities "
            "permitted to the concessionaire, rights granted to develop, "
            "operate, maintain, manage, or commercially exploit the project "
            "assets, and any explicit limitations or exclusions."
        ),
    },
    3: {
        "name": "Asset Ownership & Control",
        "query": (
            "Extract clauses specifying ownership of project assets, land, and "
            "equipment, rights of use versus title, and provisions allocating "
            "operational control, access, supervision, and decision-making authority."
        ),
    },
    4: {
        "name": "Regulatory & Operational Compliance",
        "query": (
            "Extract clauses imposing regulatory, operational, and compliance "
            "obligations, including approvals, directions, inspections, "
            "reporting requirements, standards, and the extent of supervisory "
            "or micromanagement authority over operations."
        ),
    },
    5: {
        "name": "Concession Period & Extension",
        "query": (
            "Extract clauses defining the concession term, commencement and "
            "expiry, conditions for extension or renewal, discretion of the "
            "authority, and any limits or uncertainty affecting the investment horizon."
        ),
    },
    6: {
        "name": "Tariff & Revenue Flexibility",
        "query": (
            "Extract all provisions governing how prices, tariffs, fees, or "
            "user charges are determined, revised, approved, or controlled; "
            "the extent of pricing discretion or regulation; revenue-sharing "
            "or levy mechanisms; indexation or adjustment formulas; and any "
            "restrictions affecting commercial revenue generation or innovation."
        ),
    },
    7: {
        "name": "Demand & Traffic Risk Allocation",
        "query": (
            "Extract clauses allocating demand, volume, traffic, or throughput "
            "risk, including minimum or assured traffic, exclusivity or "
            "non-exclusivity, diversion rights, competing facilities, "
            "capacity commitments, and any guarantees or disclaimers affecting traffic levels."
        ),
    },
    8: {
        "name": "Change in Law & Policy Risk",
        "query": (
            "Extract clauses addressing change in law, policy, regulation, or "
            "interpretation; allocation of resulting costs or benefits; "
            "compensation, relief, or adjustment mechanisms; and protections "
            "against adverse governmental or regulatory actions affecting the project."
        ),
    },
    9: {
        "name": "Relief Structure",
        "query": (
            "Extract clauses providing relief from contractual obligations due "
            "to force majeure, change in circumstances, or external shocks, "
            "including suspension, extension of time, cost or tariff adjustment, "
            "termination relief, and conditions or limitations on such relief."
        ),
    },
    10: {
        "name": "Termination & Step-in Rights",
        "query": (
            "Extract clauses governing termination events, defaults and cure "
            "periods, termination consequences, compensation or payout "
            "mechanisms, lender step-in or substitution rights, and protections "
            "preserving continuity of the project upon default or early termination."
        ),
    },
    11: {
        "name": "Dispute Resolution & Governing Law",
        "query": (
            "Extract clauses specifying dispute resolution mechanisms, "
            "escalation steps, arbitration or court processes, seat and venue, "
            "governing law, jurisdiction, timelines, enforceability of awards, "
            "and any provisions affecting neutrality or speed of dispute resolution."
        ),
    },
    12: {
        "name": "Assignment & Financing Flexibility",
        "query": (
            "Extract clauses governing assignment or transfer of rights and "
            "obligations, change of control, creation of security interests, "
            "financing or refinancing rights, lender protections, substitution "
            "or novation, and any approvals or restrictions affecting exit or refinancing."
        ),
    },
}


SECTION_PROMPT = ChatPromptTemplate.from_template(
    """
ROLE:
You are a senior infrastructure PPP legal–economic analyst advising
Government of India policy reform committees and multilateral lenders.

TASK:
Analyze the concession agreement ONLY for the section:
{section_name}

EVIDENCE DISCIPLINE (CRITICAL):
- Every legal conclusion MUST be backed by contract language
- Clearly distinguish between:
  (a) Explicit clauses
  (b) Implicit structure inferred from clauses
  (c) Absence or silence
- Do NOT speculate beyond the document
- If the document is silent, explicitly say so

OUTPUT FORMAT (STRICT — FOLLOW EXACTLY):

SECTION: {section_name}

LEGAL POSITION:
- What the contract explicitly allows, restricts, or mandates
- Clearly indicate if coverage is partial or conditional

ECONOMIC IMPLICATION:
- Commercial consequences strictly flowing from the contract text
- No assumptions beyond documented provisions

BANKABILITY IMPACT:
- Lender and investor perspective based on clarity, enforceability,
  and risk allocation in the contract

CLAUSE EVIDENCE:
- Verbatim excerpts from the document
- Include Article / Clause numbers wherever available
- Do NOT paraphrase in this section

EVIDENCE_STRENGTH: <STRONG | MODERATE | WEAK | NONE>
(one line only, uppercase, choose exactly one)

DOCUMENT CONTEXT:
{context}
"""
)


def upload_to_s3(file_bytes: bytes, bucket: str, key: str, region: str) -> None:
    s3 = boto3.client("s3", region_name=region)
    s3.put_object(Bucket=bucket, Key=key, Body=file_bytes)


def textract_pdf_from_s3(bucket: str, key: str, region: str) -> str:
    textract = boto3.client("textract", region_name=region)
    response = textract.start_document_text_detection(
        DocumentLocation={"S3Object": {"Bucket": bucket, "Name": key}}
    )
    job_id = response["JobId"]

    while True:
        status = textract.get_document_text_detection(JobId=job_id)
        job_status = status["JobStatus"]
        if job_status in ["SUCCEEDED", "FAILED"]:
            break
        time.sleep(5)

    if job_status != "SUCCEEDED":
        raise RuntimeError(f"Textract job failed with status: {job_status}")

    blocks: List[Dict] = []
    next_token = None

    while True:
        if next_token:
            response = textract.get_document_text_detection(
                JobId=job_id, NextToken=next_token
            )
        else:
            response = textract.get_document_text_detection(JobId=job_id)

        blocks.extend(response["Blocks"])
        next_token = response.get("NextToken")
        if not next_token:
            break

    lines = [block["Text"] for block in blocks if block["BlockType"] == "LINE"]
    return "\n".join(lines)


def save_extracted_text(text: str, document_name: str) -> Path:
    path = TEXT_DIR / f"{document_name}.txt"
    path.write_text(text, encoding="utf-8")
    return path


def build_vector_db(text: str, document_name: str) -> Path:
    vecdb = LegalVectorDB()
    vecdb.ingest_document(text, source_name=document_name)
    vecdb.build_index()
    db_path = VECTOR_DB_DIR / document_name
    vecdb.save(str(db_path))
    return db_path


def analyze_document(document_name: str, openai_key: str, top_k: int) -> List[Dict]:
    file_path = TEXT_DIR / f"{document_name}.txt"
    if not file_path.exists():
        raise FileNotFoundError(f"Missing file: {file_path}")

    vecdb_path = VECTOR_DB_DIR / document_name
    if not vecdb_path.exists():
        raise FileNotFoundError(f"Vector DB not found at: {vecdb_path}")

    vecdb = LegalVectorDB.load(vecdb_path)

    llm = ChatOpenAI(
        model="gpt-4o",
        temperature=0,
        api_key=openai_key
    )
    chain = SECTION_PROMPT | llm | StrOutputParser()

    section_out_dir = OUTPUT_DIR / document_name
    section_out_dir.mkdir(parents=True, exist_ok=True)

    results = []
    for sec_id, sec in SECTIONS.items():
        docs = vecdb.retrieve_for_section(
            section_query=sec["query"],
            top_k=top_k
        )
        context = "\n\n---\n\n".join(d.page_content for d in docs)

        result = chain.invoke({
            "section_name": sec["name"],
            "context": context
        })

        confidence = "UNKNOWN"
        match = re.search(
            r"EVIDENCE_STRENGTH:\s*(STRONG|MODERATE|WEAK|NONE)",
            result.upper()
        )
        if match:
            strength = match.group(1)
            if strength == "STRONG":
                confidence = "HIGH"
            elif strength == "MODERATE":
                confidence = "MEDIUM"
            else:
                confidence = "LOW"

        section_json = {
            "document": document_name,
            "section_id": sec_id,
            "section_name": sec["name"],
            "query_used": sec["query"],
            "analysis": result,
            "confidence": confidence
        }

        file_name = f"{sec_id:02d}_{sec['name'].replace(' ', '_')}.json"
        section_file = section_out_dir / file_name
        section_file.write_text(
            json.dumps(section_json, indent=2, ensure_ascii=False),
            encoding="utf-8"
        )

        results.append(section_json)

    return results


def generate_docx(document_name: str) -> Path:
    json_dir = OUTPUT_DIR / document_name
    if not json_dir.exists():
        raise FileNotFoundError(f"JSON output folder not found: {json_dir}")

    output_docx = BASE_DIR / "analysis_docs" / f"{document_name}_PPP_Legal_Analysis.docx"
    output_docx.parent.mkdir(exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(
        f"PPP Concession Agreement Analysis\n{document_name}",
        level=0
    )
    title.alignment = 1

    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        h.style.font.bold = True
        return h

    def add_bullet(text):
        doc.add_paragraph(text, style="List Bullet")

    def add_quote(text):
        p = doc.add_paragraph(text)
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.runs[0].italic = True

    def add_normal(text):
        doc.add_paragraph(text)

    json_files = sorted(json_dir.glob("*.json"))
    for json_file in json_files:
        data = json.loads(json_file.read_text(encoding="utf-8"))

        section_name = data["section_name"]
        analysis_md = data["analysis"]
        confidence = data.get("confidence", "UNKNOWN")

        add_heading(section_name, level=1)

        lines = analysis_md.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if line.startswith("SECTION:"):
                continue
            if line.endswith(":") and line.isupper():
                add_heading(line.replace(":", "").title(), level=2)
            elif line.startswith("- "):
                add_bullet(line[2:].strip())
            elif line.startswith("\"") or line.startswith("“"):
                add_quote(line)
            elif line.startswith("EVIDENCE_STRENGTH"):
                add_heading("Evidence Strength", level=2)
                strength = line.split(":")[1].strip()
                p = add_normal(strength)
                p = doc.paragraphs[-1]
                p.runs[0].bold = True
            else:
                add_normal(line)

        add_heading("Confidence Assessment", level=2)
        p = add_normal(f"Overall Confidence Level: {confidence}")
        p = doc.paragraphs[-1]
        p.runs[0].bold = True

        doc.add_page_break()

    doc.save(output_docx)
    return output_docx


st.set_page_config(
    page_title="PPP Legal Analysis",
    page_icon="📄",
    layout="wide"
)

st.title("PPP LLM Contract Analysis")
st.caption("Upload a PDF, extract with AWS Textract, build a vector DB, and run section-wise RAG analysis.")

load_dotenv(dotenv_path=BASE_DIR / ".env", override=True)

aws_region = os.getenv("AWS_REGION", "ap-south-1")
s3_bucket = os.getenv("S3_BUCKET", "")
s3_prefix = os.getenv("S3_PREFIX", "")
openai_key = os.getenv("OPENAI_API_KEY", "")
rag_top_k = int(os.getenv("RAG_TOP_K", "12"))

uploaded_file = st.file_uploader("Upload PDF", type=["pdf"])

document_name = st.text_input(
    "Document name",
    value=(
        uploaded_file.name.rsplit(".", 1)[0]
        if uploaded_file is not None
        else ""
    )
)

safe_name = document_name.strip()
existing_text_path = TEXT_DIR / f"{safe_name}.txt" if safe_name else None
existing_db_path = VECTOR_DB_DIR / safe_name if safe_name else None
ready_for_analysis = bool(
    safe_name
    and existing_text_path
    and existing_db_path
    and existing_text_path.exists()
    and existing_db_path.exists()
)

if "analysis_ready" not in st.session_state:
    st.session_state.analysis_ready = False

col1, col2 = st.columns(2)
with col1:
    run_extract = st.button("Extract + Build Vector DB")
with col2:
    if ready_for_analysis or st.session_state.analysis_ready:
        run_analysis = st.button("Run Analysis + Generate DOCX")
    else:
        run_analysis = False

if run_extract:
    if uploaded_file is None:
        st.error("Please upload a PDF file.")
        st.stop()
    if not s3_bucket:
        st.error("Please provide the S3 bucket name for Textract.")
        st.stop()
    if not document_name.strip():
        st.error("Please provide a document name.")
        st.stop()

    safe_name = document_name.strip()
    s3_key_base = f"{safe_name}.pdf"
    if s3_prefix.strip():
        s3_key = f"{s3_prefix.strip().rstrip('/')}/{s3_key_base}"
    else:
        s3_key = s3_key_base

    with st.status("Uploading PDF to S3...", expanded=False) as status:
        file_bytes = uploaded_file.read()
        upload_to_s3(file_bytes, s3_bucket, s3_key, aws_region)
        status.update(label="PDF uploaded to S3.", state="complete")

    with st.status("Running AWS Textract...", expanded=False) as status:
        extracted_text = textract_pdf_from_s3(s3_bucket, s3_key, aws_region)
        status.update(label="Textract extraction completed.", state="complete")

    with st.status("Saving extracted text...", expanded=False) as status:
        text_path = save_extracted_text(extracted_text, safe_name)
        status.update(label=f"Saved text to {text_path}", state="complete")

    with st.status("Building vector database...", expanded=False) as status:
        db_path = build_vector_db(extracted_text, safe_name)
        status.update(label=f"Vector DB saved to {db_path}", state="complete")

    st.session_state.analysis_ready = True

    st.success("Extraction and vector DB build completed.")
    st.subheader("Stored Artifacts")
    st.write(f"Extracted text: `{text_path}`")
    st.write(f"Vector DB: `{db_path}`")

if run_analysis:
    if not document_name.strip():
        st.error("Please provide a document name to analyze.")
        st.stop()
    if not openai_key:
        st.error("OPENAI_API_KEY is required in .env.")
        st.stop()

    safe_name = document_name.strip()
    text_path = TEXT_DIR / f"{safe_name}.txt"
    db_path = VECTOR_DB_DIR / safe_name

    if not text_path.exists():
        st.error(f"Extracted text not found: {text_path}")
        st.stop()
    if not db_path.exists():
        st.error(f"Vector DB not found: {db_path}")
        st.stop()

    with st.status("Running section-wise analysis...", expanded=False) as status:
        results = analyze_document(safe_name, openai_key, rag_top_k)
        status.update(label="Analysis completed.", state="complete")

    with st.status("Generating DOCX report...", expanded=False) as status:
        docx_path = generate_docx(safe_name)
        status.update(label=f"DOCX saved to {docx_path}", state="complete")

    st.success("Analysis and DOCX generation completed.")
    st.subheader("Outputs")
    st.write(f"JSON outputs: `{OUTPUT_DIR / safe_name}`")
    st.write(f"DOCX report: `{docx_path}`")

    for item in results:
        with st.expander(f"{item['section_id']:02d} - {item['section_name']}"):
            st.markdown(item["analysis"])
            st.write(f"Confidence: {item['confidence']}")
