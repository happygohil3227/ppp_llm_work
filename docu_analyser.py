import json
import os
import re
import sys
from pathlib import Path
from typing import Dict

from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser

from utils.vecDb import LegalVectorDB


# =====================================================
# CONFIG
# =====================================================

DOCUMENT_NAMES = [
    "Adani  Mundra Port agreement for operations",
    "Bahuli_IOCL_compressed",
    "Banthra_Kribhco_compressed",
    "CONCOR Agreement for operations_compressed",
    "GCT Pathri_Concor_compressed",
    "Kundanganj_Reliance_compressed",
    "MMLP Barhi_Concor_compressed (1)"
]

TEXT_DIR = Path("extracted_text")
OUTPUT_DIR = Path("outputs")
OUTPUT_DIR.mkdir(exist_ok=True)


# =====================================================
# 12 SECTION DEFINITIONS
# =====================================================

SECTIONS: Dict[int, Dict[str, str]] = {
    1: {"name": "Context & Objective",
        "query": "Extract the background, policy rationale, and purpose of this PPP concession agreement, including the problem it seeks to address, the public interest objective, the scope of the concession, and the intended outcomes for infrastructure service delivery and market efficiency"},
    2: {"name": "Scope of Concession & Rights Granted",
        "query": "Extract clauses defining the scope of the concession, activities permitted to the concessionaire, rights granted to develop, operate, maintain, manage, or commercially exploit the project assets, and any explicit limitations or exclusions."},
    3: {"name": "Asset Ownership & Control",
        "query": "Extract clauses specifying ownership of project assets, land, and equipment, rights of use versus title, and provisions allocating operational control, access, supervision, and decision-making authority."},
    4: {"name": "Regulatory & Operational Compliance",
        "query": "Extract clauses imposing regulatory, operational, and compliance obligations, including approvals, directions, inspections, reporting requirements, standards, and the extent of supervisory or micromanagement authority over operations."},
    5: {"name": "Concession Period & Extension",
        "query": "Extract clauses defining the concession term, commencement and expiry, conditions for extension or renewal, discretion of the authority, and any limits or uncertainty affecting the investment horizon."},
    6: {"name": "Tariff & Revenue Flexibility",
        "query": "Extract all provisions governing how prices, tariffs, fees, or user charges are determined, revised, approved, or controlled; the extent of pricing discretion or regulation; revenue-sharing or levy mechanisms; indexation or adjustment formulas; and any restrictions affecting commercial revenue generation or innovation."},
    7: {"name": "Demand & Traffic Risk Allocation",
        "query": "Extract clauses allocating demand, volume, traffic, or throughput risk, including minimum or assured traffic, exclusivity or non-exclusivity, diversion rights, competing facilities, capacity commitments, and any guarantees or disclaimers affecting traffic levels."},
    8: {"name": "Change in Law & Policy Risk",
        "query": "Extract clauses addressing change in law, policy, regulation, or interpretation; allocation of resulting costs or benefits; compensation, relief, or adjustment mechanisms; and protections against adverse governmental or regulatory actions affecting the project."},
    9: {"name": "Relief Structure",
        "query": "Extract clauses providing relief from contractual obligations due to force majeure, change in circumstances, or external shocks, including suspension, extension of time, cost or tariff adjustment, termination relief, and conditions or limitations on such relief."},
    10: {"name": "Termination & Step-in Rights",
         "query": "Extract clauses governing termination events, defaults and cure periods, termination consequences, compensation or payout mechanisms, lender step-in or substitution rights, and protections preserving continuity of the project upon default or early termination."},
    11: {"name": "Dispute Resolution & Governing Law",
         "query": "Extract clauses specifying dispute resolution mechanisms, escalation steps, arbitration or court processes, seat and venue, governing law, jurisdiction, timelines, enforceability of awards, and any provisions affecting neutrality or speed of dispute resolution."},
    12: {"name": "Assignment & Financing Flexibility",
         "query": "Extract clauses governing assignment or transfer of rights and obligations, change of control, creation of security interests, financing or refinancing rights, lender protections, substitution or novation, and any approvals or restrictions affecting exit or refinancing."}
}


# =====================================================
# SECTION EXTRACTION PROMPT
# =====================================================

SECTION_PROMPT = ChatPromptTemplate.from_template("""
ROLE:
You are a senior infrastructure PPP legal–economic analyst advising
Government of India policy reform committees and multilateral lenders.

TASK:
Analyze the concession agreement ONLY for the section:
{section_name}

EVIDENCE DISCIPLINE (CRITICAL):
- Every legal conclusion MUST be backed by contract language
- Clearly distinguish between:
  (a) Explicit clauses
  (b) Implicit structure inferred from clauses
  (c) Absence or silence
- Do NOT speculate beyond the document
- If the document is silent, explicitly say so

OUTPUT FORMAT (STRICT — FOLLOW EXACTLY):

SECTION: {section_name}

LEGAL POSITION:
- What the contract explicitly allows, restricts, or mandates
- Clearly indicate if coverage is partial or conditional

ECONOMIC IMPLICATION:
- Commercial consequences strictly flowing from the contract text
- No assumptions beyond documented provisions

BANKABILITY IMPACT:
- Lender and investor perspective based on clarity, enforceability,
  and risk allocation in the contract

CLAUSE EVIDENCE:
- Verbatim excerpts from the document
- Include Article / Clause numbers wherever available
- Do NOT paraphrase in this section

EVIDENCE_STRENGTH: <STRONG | MODERATE | WEAK | NONE>
(one line only, uppercase, choose exactly one)

DOCUMENT CONTEXT:
{context}
""")


def analyze_document(document_name: str):
    print(f"\n📄 Processing document: {document_name}")

    file_path = TEXT_DIR / f"{document_name}.txt"
    if not file_path.exists():
        raise FileNotFoundError(f"Missing file: {file_path}")

    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()

    vector_db_dir = Path("vectorDB") / document_name
    vecdb = LegalVectorDB.load(vector_db_dir)

    llm = ChatOpenAI(
        model="gpt-4o",
        temperature=0,
        api_key=os.getenv("OPENAI_API_KEY")
    )

    chain = SECTION_PROMPT | llm | StrOutputParser()

    section_out_dir = OUTPUT_DIR / document_name
    section_out_dir.mkdir(parents=True, exist_ok=True)

    for sec_id, sec in SECTIONS.items():
        print(f"   ➤ Section {sec_id}: {sec['name']}")

        docs = vecdb.retrieve_for_section(
            section_query=sec["query"],
            top_k=12
        )

        context = "\n\n---\n\n".join(d.page_content for d in docs)

        result = chain.invoke({
            "section_name": sec["name"],
            "context": context
        })

        confidence = "UNKNOWN"
        match = re.search(r"EVIDENCE_STRENGTH:\s*(STRONG|MODERATE|WEAK|NONE)", result.upper())
        if match:
            strength = match.group(1)
            if strength == "STRONG":
                confidence = "HIGH"
            elif strength == "MODERATE":
                confidence = "MEDIUM"
            elif strength in ["WEAK", "NONE"]:
                confidence = "LOW"

        section_json = {
            "document": document_name,
            "section_id": sec_id,
            "section_name": sec["name"],
            "query_used": sec["query"],
            "analysis": result,
            "confidence": confidence
        }
        print(result)

        file_name = f"{sec_id:02d}_{sec['name'].replace(' ', '_')}.json"
        section_file = section_out_dir / file_name

        with open(section_file, "w", encoding="utf-8") as f:
            json.dump(section_json, f, indent=2, ensure_ascii=False)

        print(f"      ✓ Saved {file_name}")

    print(f"\n✅ Completed analysis for document: {document_name}")


def generate_docx(document_name: str) -> Path:
    json_dir = Path("outputs") / document_name
    output_docx = Path(f"analysis_docs/{document_name}_PPP_Legal_Analysis.docx")
    output_docx.parent.mkdir(exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(
        f"PPP Concession Agreement Analysis\n{document_name}",
        level=0
    )
    title.alignment = 1

    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        h.style.font.bold = True
        return h

    def add_bullet(text):
        doc.add_paragraph(text, style="List Bullet")

    def add_quote(text):
        p = doc.add_paragraph(text)
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.runs[0].italic = True

    def add_normal(text):
        doc.add_paragraph(text)

    json_files = sorted(json_dir.glob("*.json"))
    for json_file in json_files:
        with open(json_file, "r", encoding="utf-8") as f:
            data = json.load(f)

        section_name = data["section_name"]
        analysis_md = data["analysis"]
        confidence = data.get("confidence", "UNKNOWN")

        add_heading(section_name, level=1)

        lines = analysis_md.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue

            if line.startswith("SECTION:"):
                continue

            if line.endswith(":") and line.isupper():
                add_heading(line.replace(":", "").title(), level=2)
            elif line.startswith("- "):
                add_bullet(line[2:].strip())
            elif line.startswith("\"") or line.startswith("“"):
                add_quote(line)
            elif line.startswith("EVIDENCE_STRENGTH"):
                add_heading("Evidence Strength", level=2)
                strength = line.split(":")[1].strip()
                p = add_normal(strength)
                p = doc.paragraphs[-1]
                p.runs[0].bold = True
            else:
                add_normal(line)

        add_heading("Confidence Assessment", level=2)
        p = add_normal(f"Overall Confidence Level: {confidence}")
        p = doc.paragraphs[-1]
        p.runs[0].bold = True

        doc.add_page_break()

    doc.save(output_docx)
    return output_docx


def main() -> None:
    if len(sys.argv) < 2:
        raise SystemExit("Usage: python docu_analyser.py <DOCUMENT_NAME>")

    load_dotenv(dotenv_path=Path(".env"), override=True)
    doc_name = sys.argv[1]

    analyze_document(doc_name)
    output_docx = generate_docx(doc_name)
    print(f"✅ PPP Legal Memo generated: {output_docx.resolve()}")


if __name__ == "__main__":
    main()
