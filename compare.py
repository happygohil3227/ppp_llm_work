import json
import os
import sys
from pathlib import Path
from typing import Dict, List

from dotenv import load_dotenv
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_openai import ChatOpenAI
from docx import Document
from docx.shared import Pt


# =========================
# CONFIG
# =========================

OUTPUTS_DIR = Path("outputs")
COMPARE_ROOT = Path("doc_compare")

QUERIES = [
    {
        "id": 1,
        "title": "Scope of Concession & Rights Granted",
        "query": (
            "Which MCA grants greater commercial autonomy to the Concessionaire through "
            "the scope of concession and rights granted, and how does this affect the "
            "ability to undertake ancillary activities, optimize operations, and "
            "generate additional revenue?\n\n"
            "Compare both agreements using explicit rights, implied limitations, and "
            "material silence, supported by clause evidence and evidence strength assessment."
        )
    },
    {
        "id": 2,
        "title": "Tariff & Revenue Flexibility",
        "query": (
            "Which MCA provides greater flexibility in tariff setting and revenue "
            "generation for the Concessionaire, and how does this impact cash flow "
            "certainty and lender comfort?\n\n"
            "Compare both agreements based on explicit pricing clauses, regulatory "
            "controls, implicit constraints, and contractual silence, with clause "
            "evidence and evidence strength assessment."
        )
    },
    {
        "id": 3,
        "title": "Asset Ownership & Control",
        "query": (
            "Which MCA allocates asset ownership and operational control more favorably "
            "to the Concessionaire, and how does this affect control over operations, "
            "monetization potential, and lender security?\n\n"
            "Compare both agreements using ownership clauses, control rights, "
            "restrictions, and clause evidence with evidence strength assessment."
        )
    },
    {
        "id": 4,
        "title": "Demand & Traffic Risk",
        "query": (
            "Which MCA allocates demand or traffic risk more efficiently between the "
            "Authority and the Concessionaire, and how does this influence revenue "
            "stability, downside risk, and bankability?\n\n"
            "Compare both agreements based on explicit risk allocation, guarantees, "
            "exclusivity provisions, and material silence, supported by clause evidence "
            "and evidence strength assessment."
        )
    },
    {
        "id": 5,
        "title": "Regulatory Micromanagement",
        "query": (
            "Which MCA subjects the Concessionaire to greater regulatory or operational "
            "micromanagement, and how does this affect operational autonomy, compliance "
            "costs, and efficiency?\n\n"
            "Compare both agreements using explicit approval requirements, supervisory "
            "powers, implied control mechanisms, and clause evidence with evidence strength assessment."
        )
    },
    {
        "id": 6,
        "title": "Dispute Resolution",
        "query": (
            "Which MCA provides a more neutral, predictable, and enforceable dispute "
            "resolution framework, and how does this affect legal certainty and lender "
            "confidence?\n\n"
            "Compare both agreements based on arbitration structure, appointment "
            "mechanisms, governing law, enforcement certainty, and clause evidence with evidence strength assessment."
        )
    },
    {
        "id": 7,
        "title": "Termination",
        "query": (
            "Which MCA exposes the Concessionaire to higher termination risk, and how "
            "does this affect equity protection, cash flow continuity, and lender "
            "confidence?\n\n"
            "Compare both agreements using termination triggers, cure periods, "
            "compensation provisions, and clause evidence with evidence strength assessment."
        )
    },
    {
        "id": 8,
        "title": "Exit & End-of-Term Provisions",
        "query": (
            "Which MCA provides greater flexibility and certainty for exit and "
            "end-of-term outcomes for the Concessionaire, and how does this affect "
            "recoverability of investment and refinancing options?\n\n"
            "Compare both agreements based on exit rights, transfer provisions, "
            "handover obligations, residual value treatment, and clause evidence with evidence strength assessment."
        )
    }
]


# =========================
# HELPERS
# =========================

def load_section_json(doc_name: str) -> Dict[int, Dict]:
    doc_dir = OUTPUTS_DIR / doc_name
    if not doc_dir.exists():
        raise FileNotFoundError(f"Missing outputs folder: {doc_dir}")

    data = {}
    for path in sorted(doc_dir.glob("*.json")):
        payload = json.loads(path.read_text(encoding="utf-8"))
        data[payload["section_id"]] = payload
    return data


def make_llm() -> ChatOpenAI:
    return ChatOpenAI(model="gpt-4o", temperature=0, api_key=os.getenv("OPENAI_API_KEY"))


# =========================
# PROMPTS
# =========================

AGENT1_PROMPT = ChatPromptTemplate.from_template('''
ROLE:
You are Agent 1 in a two-agent comparison pipeline.

TASK:
Given the user query, choose the minimum necessary sections from the 12-section PPP framework.
Return ONLY the section IDs as a comma-separated list (e.g., "6,7,8").

USER QUERY:
{user_query}

12 SECTION NAMES:
1. Context & Objective
2. Scope of Concession & Rights Granted
3. Asset Ownership & Control
4. Regulatory & Operational Compliance
5. Concession Period & Extension
6. Tariff & Revenue Flexibility
7. Demand & Traffic Risk Allocation
8. Change in Law & Policy Risk
9. Relief Structure
10. Termination & Step-in Rights
11. Dispute Resolution & Governing Law
12. Assignment & Financing Flexibility
''')


AGENT2_PROMPT = ChatPromptTemplate.from_template('''
ROLE:
You are Agent 2 in a two-agent comparison pipeline.

TASK:
Use the extracted section analyses from both documents to answer the user query.
Compare the two documents and produce a structured output exactly in this format:

Finding:
...

Comparison Insight:
...

Impact on Cash Flows:
...

Clause Evidence:
- Document A: <verbatim excerpts with clause/article numbers>
- Document B: <verbatim excerpts with clause/article numbers>

Evidence Strength: <HIGH | MEDIUM | LOW>

RULES:
- Base findings ONLY on the provided section analyses.
- Use clause evidence from those analyses.
- Do not speculate beyond the text.
- If evidence is missing, say so explicitly.

USER QUERY:
{user_query}

DOCUMENT A NAME:
{doc_a_name}

DOCUMENT B NAME:
{doc_b_name}

DOCUMENT A SECTION ANALYSES:
{doc_a_sections}

DOCUMENT B SECTION ANALYSES:
{doc_b_sections}
''')


def run_compare(doc_a: str, doc_b: str) -> Path:
    compare_dir = COMPARE_ROOT / f"{doc_a}_&{doc_b}_compare"
    compare_dir.mkdir(parents=True, exist_ok=True)

    llm = make_llm()
    agent1_chain = AGENT1_PROMPT | llm | StrOutputParser()
    agent2_chain = AGENT2_PROMPT | llm | StrOutputParser()

    doc_a_sections = load_section_json(doc_a)
    doc_b_sections = load_section_json(doc_b)

    results = []

    for q in QUERIES:
        print(f"Processing query {q['id']:02d}: {q['title']}")
        section_ids_raw = agent1_chain.invoke({"user_query": q["query"]})
        section_ids = [int(s.strip()) for s in section_ids_raw.split(",") if s.strip().isdigit()]

        selected_a = [doc_a_sections[sid] for sid in section_ids if sid in doc_a_sections]
        selected_b = [doc_b_sections[sid] for sid in section_ids if sid in doc_b_sections]

        payload_a = "\n\n---\n\n".join([s["analysis"] for s in selected_a])
        payload_b = "\n\n---\n\n".join([s["analysis"] for s in selected_b])

        final_answer = agent2_chain.invoke({
            "user_query": q["query"],
            "doc_a_name": doc_a,
            "doc_b_name": doc_b,
            "doc_a_sections": payload_a,
            "doc_b_sections": payload_b,
        })

        out_json = {
            "query_id": q["id"],
            "title": q["title"],
            "query": q["query"],
            "section_ids": section_ids,
            "analysis": final_answer
        }

        file_name = f"{q['id']:02d}_{q['title'].replace(' ', '_')}.json"
        out_path = compare_dir / file_name
        out_path.write_text(json.dumps(out_json, indent=2, ensure_ascii=False), encoding="utf-8")

        results.append(out_json)
        print(f"  ✓ Saved query output: {out_path}")

    print(f"Saved {len(results)} query outputs to: {compare_dir}")
    return compare_dir


def generate_compare_docx(doc_a: str, doc_b: str, compare_dir: Path) -> Path:
    output_docx = COMPARE_ROOT / f"{doc_a}_&{doc_b}_compare.docx"

    json_files = sorted(compare_dir.glob("*.json"))

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(
        f"MCA Comparison Analysis {doc_a} vs {doc_b}",
        level=0
    )
    title.alignment = 1

    for json_file in json_files:
        data = json.loads(json_file.read_text(encoding="utf-8"))

        doc.add_heading(f"{data['query_id']:02d}. {data['title']}", level=1)
        doc.add_paragraph(data["query"], style="Intense Quote")

        lines = data["analysis"].split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue

            if line.endswith(":") and line.replace(":", "").istitle():
                doc.add_heading(line.replace(":", ""), level=2)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:].strip(), style="List Bullet")
            else:
                doc.add_paragraph(line)

        doc.add_page_break()

    doc.save(output_docx)
    print(f"  ✓ DOCX saved: {output_docx}")
    return output_docx


def main() -> None:
    if len(sys.argv) < 3:
        raise SystemExit("Usage: python compare.py <DOC_A> <DOC_B>")

    load_dotenv(dotenv_path=Path(".env"), override=True)

    doc_a = sys.argv[1]
    doc_b = sys.argv[2]

    compare_dir = run_compare(doc_a, doc_b)
    output_docx = generate_compare_docx(doc_a, doc_b, compare_dir)
    print(f"DOCX saved: {output_docx}")


if __name__ == "__main__":
    main()
